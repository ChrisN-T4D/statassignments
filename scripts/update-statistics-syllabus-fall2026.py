"""Build PSYC 4213 Statistics Fall 2026 syllabus (Methods Market + NWOSU calendar)."""
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r"C:\Users\couga\Downloads\Psych 4123- Neu - Statistics - Spring 2026.docx"
OUT = r"C:\Users\couga\Downloads\Psych 4213 - Statistics - Fall 2026.docx"
OUT_INPERSON = r"C:\Users\couga\Downloads\Psych 4213 - Statistics - Fall 2026 (In Person).docx"


def replace_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph(doc, startswith):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    return None


def find_paragraph_contains(doc, fragment):
    for p in doc.paragraphs:
        if fragment in p.text:
            return p
    return None


def replace_section_after_header(doc, header, content):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(header):
            if i + 1 < len(doc.paragraphs):
                replace_paragraph_text(doc.paragraphs[i + 1], content)
            return
    p = find_paragraph_contains(doc, content[:40])
    if p:
        replace_paragraph_text(p, content)


def set_table_rows(table, rows):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val


COURSE_MATERIALS = (
    "We use free, open-access materials — no textbook purchase required.\n\n"
    "Primary textbook: Learning Statistics with jamovi (LSJ), read online at "
    "https://davidfoxcroft.github.io/lsj-book/\n\n"
    "Course platform: Canvas (https://nwosu.instructure.com) holds the syllabus, weekly modules, "
    "due dates, grades, discussions, and graded benchmarks.\n\n"
    "Practice platform: Methods Market (https://methods-market.clneu.com/class/statistics) holds "
    "Concept Review quizzes and Software Practice for each module. Due dates and points are always "
    "in Canvas — not Methods Market.\n\n"
    "Methods Market Setup (complete by end of Week 1)\n"
    "Step 1 — Create your account at https://methods-market.clneu.com/auth\n"
    "Step 2 — Link your student key at https://methods-market.clneu.com/claim (semester Fall 2026 / "
    "2026FA; key from instructor)\n"
    "Step 3 — Bookmark https://methods-market.clneu.com/class/statistics and Assignment Help "
    "(…/class/statistics/assignment-help)\n"
    "Step 4 — Install jamovi before Module 3: https://www.jamovi.org/download.html\n"
    "Each week: read the LSJ chapter → Concept Review → Software Practice (Modules 3–8) → complete "
    "Canvas assignments. See also the Canvas page Getting Started on Methods Market.\n\n"
    "Supplementary materials may include instructor videos, open-access articles, and other web "
    "resources linked from Canvas modules."
)

PERFORMANCE_ASSESSMENT = (
    "The readings, instructor videos, and Methods Market activities present the major topics in "
    "statistics along with recent developments in the field. To build a thorough understanding, "
    "complete the assigned LSJ chapter reading, finish Concept Review for that module, and complete "
    "Software Practice (Modules 3–8) before graded Jamovi assignments and benchmarks. Readings and "
    "videos often contain unique material (presented in one medium but not the other).\n\n"
    "Recommended preparation for benchmarks: read the assigned chapters, complete Concept Review "
    "(and Software Practice where assigned), work through Methods Market topics, then use the "
    "formative benchmark practice in Methods Market before the proctored Canvas benchmark. Post "
    "questions in the module discussion board before emailing. Talking with classmates and attending "
    "office hours helps organize concepts.\n\n"
    "This is an ideas-and-application course. Assessments stress concepts, relationships among "
    "concepts, and application to practical examples — not isolated memorization."
)

BENCHMARKS = (
    "There will be three benchmarks during the semester. Your benchmark average accounts for 50% of "
    "your final grade, weighted as follows within that category: Benchmark 1 = 15%, Benchmark 2 = "
    "15%, Benchmark 3 (final exam) = 20%.\n\n"
    "Benchmark 1 comes after Module 3 (Chapters 1–3), Benchmark 2 after Module 5 (Methods Market "
    "Modules 4–5; LSJ Chapters 4–6), and Benchmark 3 during NWOSU final exam week Dec 7–11 "
    "(Modules 6–8 only; LSJ Chapters 7–13; not a comprehensive final). Benchmarks may include multiple-choice, essay, and numerical "
    "questions. All three are proctored per university policy using Respondus LockDown Browser with "
    "webcam in Canvas.\n\n"
    "Before each graded benchmark, complete Concept Review for the covered modules in Methods Market. "
    "Optional formative practice tests (Assignment Help → Benchmark practice links) help identify "
    "weak areas.\n\n"
    "There are no make-up benchmarks for illness, approved school projects, religious holidays, "
    "family emergencies, or anything else (arrangements will be made for approved student-athlete "
    "activities). A missed benchmark counts as zero. Benchmark responses are checked for plagiarism.\n\n"
    "If you believe a benchmark question was ambiguous, you have one week from the date taken to "
    "raise the issue with the instructor. No exceptions to this deadline."
)

ASSIGNMENTS = (
    "Your assignments average accounts for 35% of your final grade, weighted as follows within that "
    "category: Methods Market Concept Review (Modules 1–8) = 10%; Methods Market Software Practice "
    "(Modules 3–8) = 5%; Jamovi screen-recording and data-analysis assignments = 15%; Canvas quizzes "
    "(e.g., Levels of Measurement, Normality) = 5%.\n\n"
    "Details and due dates are in each Canvas module. Tips and Methods Market links: "
    "https://methods-market.clneu.com/class/statistics/assignment-help\n\n"
    "Due dates still matter for pacing and feedback. If you need more time, contact me. Late coursework "
    "is accepted without an automatic grade penalty unless a specific item (for example a locked, "
    "proctored benchmark window) says otherwise. Assignments are checked for plagiarism."
)

DISCUSSIONS = (
    "Your discussion average accounts for 15% of your final grade, weighted as follows within that "
    "category: Module Questions & Discussion (ongoing Q&A each module) = 5%; graded discussion "
    "prompts (Modules 1, 2, 4, 7, and 8) = 10%.\n\n"
    "Each module includes a Module Questions & Discussion board for course questions (reading, "
    "Concept Review, software, assignments). Post here before emailing when possible. Graded prompts "
    "include AI-comparison exercises and applied examples; follow Canvas instructions for initial "
    "posts and replies.\n\n"
    "Due dates still matter for pacing and feedback. If you need more time, contact me. Late discussions "
    "or responses are accepted without an automatic grade penalty unless a specific item says otherwise. "
    "Discussion posts are checked for plagiarism."
)

EVALUATION_AND_GRADING = (
    "EVALUATION AND GRADING\n\n"
    "Final course grade components (total 100%):\n\n"
    "Benchmarks — 50% of final grade\n"
    "  • Benchmark 1 (Modules 1–3, Ch. 1–3) — 15%\n"
    "  • Benchmark 2 (Modules 4–5, Ch. 4–6) — 15%\n"
    "  • Benchmark 3 / Final Exam (Modules 6–8 only, Ch. 7–13; not a comprehensive final) — 20%\n\n"
    "Assignments — 35% of final grade\n"
    "  • Methods Market Concept Review (Modules 1–8) — 10%\n"
    "  • Methods Market Software Practice (Modules 3–8) — 5%\n"
    "  • Jamovi / data-analysis assignments & screen recordings — 15%\n"
    "  • Canvas quizzes — 5%\n\n"
    "Discussions — 15% of final grade\n"
    "  • Module Questions & Discussion (Q&A) — 5%\n"
    "  • Graded discussion prompts (Modules 1, 2, 4, 7, 8) — 10%\n\n"
    "Final Course Grade:\n"
    "A = 90–100\n"
    "B = 80–89\n"
    "C = 70–79\n"
    "D = 60–69\n"
    "F = 0–59\n\n"
    "Extra Credit: There are no extra credit opportunities. Make each assignment count.\n\n"
    "Canvas assignment groups mirror these categories. Point values for individual assignments are "
    "listed in Canvas; category weights above determine your final grade."
)

TESTING = (
    "Benchmarks 1, 2, and 3 are completed in Canvas. Benchmark 3 is the final exam during NWOSU "
    "final exam week (Dec 7–11, 2026). It covers Modules 6–8 only (not a comprehensive final). "
    "Proctored testing is required.\n\n"
    "OPTION 1 — On-campus proctoring (Alva, Woodward, Enid, Ponca City): Make prior arrangements. "
    "No walk-in testing. There may be an additional fee depending on location — check early.\n\n"
    "OPTION 2 — Respondus LockDown Browser with webcam on your personal computer. Install from the "
    "link in your Canvas course. See https://www.nwosu.instructure.com for online testing tips.\n\n"
    "Benchmark 1: proctored, due end of Week 6 (Sep 27, 2026).\n"
    "Benchmark 2: proctored, due end of Week 11 (Nov 1, 2026).\n"
    "Benchmark 3 / Final: proctored, Dec 7–11, 2026 (semester ends Dec 11)."
)

# NWOSU Fall 2026 academic calendar: classes begin Aug 17; Labor Day Sep 7; fall break Oct 15–16;
# Thanksgiving Nov 25–27; finals Dec 7–11.
SCHEDULE_ROWS = [
    ("1", "Aug 17 – Aug 23", "Module 1", "Course Introduction & Why Learn Statistics — read Ch. 1; Concept Review; Methods Market setup", "Ch. 1"),
    ("2", "Aug 24 – Aug 30", "Module 2", "Research Design & Measurement — read Ch. 2; Concept Review", "Ch. 2"),
    ("3", "Aug 31 – Sep 06", "Module 2", "Research Design & Measurement (cont.) — PSS assignment; discussion", "Ch. 2"),
    ("4", "Sep 07 – Sep 13", "Module 3", "Introduction to jamovi & Data Handling — read Ch. 3; install jamovi (Labor Day Mon Sep 7)", "Ch. 3"),
    ("5", "Sep 14 – Sep 20", "Module 3", "jamovi & Data Handling (cont.) — Software Practice; Jamovi assignments", "Ch. 3"),
    ("6", "Sep 21 – Sep 27", "—", "Benchmark 1 (Proctored, LockDown Browser) — Review Ch. 1–3", "Ch. 1–3"),
    ("7", "Sep 28 – Oct 04", "Module 4", "Descriptive Statistics — read Ch. 4; Concept Review & Software Practice", "Ch. 4"),
    ("8", "Oct 05 – Oct 11", "Module 4", "Descriptive Statistics (cont.) — Jamovi assignment; discussion", "Ch. 4"),
    ("9", "Oct 12 – Oct 18", "Module 5", "Graphing & Visualization — read Ch. 5–6; Concept Review & Software Practice (Fall Break Oct 15–16)", "Ch. 5, 6"),
    ("10", "Oct 19 – Oct 25", "Module 5", "Graphing & Visualization (cont.) — screen-record assignment", "Ch. 5, 6"),
    ("11", "Oct 26 – Nov 01", "—", "Benchmark 2 (Proctored, LockDown Browser) — Review Modules 4–5 (Ch. 4–6)", "Ch. 4–6"),
    ("12", "Nov 02 – Nov 08", "Module 6", "Probability & Sampling — read Ch. 7–8; Concept Review & Software Practice", "Ch. 7, 8"),
    ("13", "Nov 09 – Nov 15", "Module 6", "Probability & Sampling (cont.) — dice/normal/CLT assignments", "Ch. 7, 8"),
    ("14", "Nov 16 – Nov 22", "Module 7", "Hypothesis Testing — read Ch. 9; Concept Review & Software Practice", "Ch. 9"),
    ("15", "Nov 23 – Nov 29", "Module 7", "Hypothesis Testing (cont.) — hypotheses assignment; discussion (Thanksgiving Nov 25–27)", "Ch. 9"),
    ("16", "Nov 30 – Dec 06", "Module 8", "Comparing Groups & Relationships — Ch. 10–13; Week 16 assignment", "Ch. 10–13"),
    ("Finals", "Dec 07 – Dec 11", "—", "Benchmark 3 / Final Exam (Proctored; Modules 6–8 only, not comprehensive; Ch. 7–13)", "Ch. 7–13"),
]

SCHEDULE_FOOTNOTE = (
    "Calendar notes (NWOSU Fall 2026): Regular classes begin August 17. Labor Day: September 7. "
    "Fall Break: October 15–16. Thanksgiving Break: November 25–27 (no classes). Final examinations: "
    "December 7–11."
)

ATTENDANCE_ONLINE = (
    "As this class is online, we do not meet in person. Module discussion boards are where you ask "
    "questions, connect with classmates, and complete graded discussion prompts. Use them actively — "
    "they count as our shared class meeting time."
)

ATTENDANCE_INPERSON = (
    "This section meets on campus. Regular attendance is expected per university policy. "
    "Module discussion boards are for questions between class meetings and for graded discussion "
    "prompts — post there before emailing when possible."
)

MEETING_ONLINE = "Online — asynchronous (no scheduled meeting time)"

MEETING_INPERSON = (
    "On campus — see your section listing in RangerNet / the official course schedule for meeting "
    "days, times, and room (sections A025, E926, W094, X098, X150)."
)


def main():
    in_person = "--in-person" in sys.argv
    doc = Document(SRC)

    replace_paragraph_text(doc.paragraphs[1], "Psych 4213 – Statistics")
    replace_paragraph_text(doc.paragraphs[2], "Fall 2026")
    replace_paragraph_text(doc.paragraphs[16], "In Person" if in_person else "Online")

    # Meeting Times and Location (paragraph after Course Modality value)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Meeting Times and Location") and i + 1 < len(doc.paragraphs):
            replace_paragraph_text(
                doc.paragraphs[i + 1],
                MEETING_INPERSON if in_person else MEETING_ONLINE,
            )
            break
        # Sometimes the value is the next non-empty paragraph after a bare header line
        if p.text.strip() == "Meeting Times and Location:" and i + 1 < len(doc.paragraphs):
            replace_paragraph_text(
                doc.paragraphs[i + 1],
                MEETING_INPERSON if in_person else MEETING_ONLINE,
            )
            break

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("We use free") or p.text.strip().startswith("We will use free"):
            replace_paragraph_text(p, COURSE_MATERIALS)
            break

    replace_section_after_header(doc, "Performance Assessment:", PERFORMANCE_ASSESSMENT)
    replace_section_after_header(doc, "Benchmarks:", BENCHMARKS)
    replace_section_after_header(doc, "Assignments:", ASSIGNMENTS)
    replace_section_after_header(doc, "Discussions:", DISCUSSIONS)

    for p in doc.paragraphs:
        if p.text.startswith("Students are responsible for regularly checking both"):
            replace_paragraph_text(
                p,
                "Students are responsible for regularly checking both Canvas and their NWOSU email "
                "during all university closures. These include closures caused by weather, safety "
                "concerns, or other emergencies. This course will adapt as necessary to support "
                "continued progress toward achieving learning outcomes. Instructors will provide "
                "updates, expectations, and instructions through NWOSU email and Canvas."
            )
            break

    eval_p = find_paragraph(doc, "EVALUATION AND GRADING")
    if eval_p:
        replace_paragraph_text(eval_p, EVALUATION_AND_GRADING)
        start = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith("EVALUATION AND GRADING"):
                start = i
                break
        if start is not None:
            for j in range(start + 1, start + 15):
                if j >= len(doc.paragraphs):
                    break
                t = doc.paragraphs[j].text.strip()
                if t.startswith("Student Complaint") or t.startswith("Class Attendance"):
                    break
                if t in ("Example:", "Benchmarks: 50% of total grade", "Assignments: 35% of total grade",
                         "Discussions: 15% of total grade", "Final Course Grade:", "A= 90-100",
                         "B= 80-89", "C= 70-79", "D= 60-69", "F= 0-59") or t.startswith("Extra Credit"):
                    replace_paragraph_text(doc.paragraphs[j], "")

    # Keep Testing identical for online and in-person (LockDown options); only modality /
    # meeting location / attendance differ between the two Fall syllabi.
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Testing and Proctoring"):
            replace_paragraph_text(p, TESTING)
            for j in range(i + 1, i + 10):
                if j >= len(doc.paragraphs):
                    break
                t = doc.paragraphs[j].text.strip()
                if t.startswith("Course Outline"):
                    break
                if t:
                    replace_paragraph_text(doc.paragraphs[j], "")
            break

    outline_p = find_paragraph(doc, "Course Outline and Tentative Schedule")
    if outline_p:
        insert_paragraph_after(outline_p, SCHEDULE_FOOTNOTE)

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Final Exam Date and Time") and i + 1 < len(doc.paragraphs):
            replace_paragraph_text(
                doc.paragraphs[i + 1],
                "Benchmark 3 (Final Exam): Dec 7–11, 2026; Modules 6–8 only (not a comprehensive final). Proctored, LockDown Browser + Webcam. Semester ends Dec 11."
            )
            break

    attendance_p = find_paragraph(doc, "Class Attendance")
    if attendance_p:
        # Header may be "Class Attendance:" with body in same or next paragraph
        body = ATTENDANCE_INPERSON if in_person else ATTENDANCE_ONLINE
        if attendance_p.text.strip() in ("Class Attendance", "Class Attendance:"):
            if attendance_p.text.strip().endswith(":"):
                # body often in next paragraph
                idx = doc.paragraphs.index(attendance_p)
                if idx + 1 < len(doc.paragraphs):
                    replace_paragraph_text(doc.paragraphs[idx + 1], body)
                else:
                    replace_paragraph_text(attendance_p, "Class Attendance:\n" + body)
            else:
                replace_paragraph_text(attendance_p, "Class Attendance:\n" + body)
        else:
            replace_paragraph_text(attendance_p, body)

    for p in doc.paragraphs:
        if "Blackboard" in p.text or "blackboard" in p.text:
            replace_paragraph_text(p, p.text.replace("Blackboard", "Canvas").replace("blackboard", "Canvas"))

    set_table_rows(doc.tables[0], SCHEDULE_ROWS)

    doc.save(OUT_INPERSON if in_person else OUT)
    print(f"Saved: {OUT_INPERSON if in_person else OUT}")


if __name__ == "__main__":
    main()
