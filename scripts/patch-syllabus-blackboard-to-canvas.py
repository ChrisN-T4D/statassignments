"""Replace Blackboard / BlackBoard / blackboard with Canvas in syllabus DOCX files."""
from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn


def replace_paragraph_text(paragraph, new_text):
    """Replace paragraph text, including hyperlink runs Word stores outside paragraph.runs."""
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:pPr"):
            continue
        p.remove(child)
    paragraph.add_run(new_text)


def fix_text(t: str) -> str:
    # Fix mangled double URLs from prior half-replacements
    t = t.replace(
        "https://nwosu.instructure.comhttps://www.nwosu.edu/blackboard",
        "https://nwosu.instructure.com",
    )
    t = t.replace(
        "https://www.nwosu.edu/blackboardhttps://nwosu.instructure.com",
        "https://nwosu.instructure.com",
    )
    t = t.replace(
        "https://nwosu.instructure.comhttps://www.nwosu.edu/Blackboard",
        "https://nwosu.instructure.com",
    )
    t = t.replace(
        "https://www.nwosu.edu/Blackboardhttps://nwosu.instructure.com",
        "https://nwosu.instructure.com",
    )
    t = t.replace("https://www.nwosu.edu/blackboard", "https://nwosu.instructure.com")
    t = t.replace("https://www.nwosu.edu/Blackboard", "https://nwosu.instructure.com")
    t = t.replace("BlackBoard", "Canvas")
    t = t.replace("Blackboard", "Canvas")
    t = t.replace("blackboard", "Canvas")
    return t


def main():
    downloads = Path(r"C:\Users\couga\Downloads")
    files = []
    for pat in ("*Psych*.docx", "*PSYC*.docx", "*Syllabus*.docx"):
        files.extend(downloads.glob(pat))

    seen = set()
    changed = []
    for f in sorted(files, key=lambda p: p.name.lower()):
        r = f.resolve()
        if r in seen or "module_" in f.name.lower():
            continue
        seen.add(r)
        try:
            doc = Document(str(f))
        except Exception as e:
            print("SKIP", f.name, e)
            continue
        n = 0
        for p in doc.paragraphs:
            if re.search(r"[Bb]lack[Bb]oard", p.text):
                new = fix_text(p.text)
                if new != p.text:
                    replace_paragraph_text(p, new)
                    n += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if re.search(r"[Bb]lack[Bb]oard", p.text):
                            new = fix_text(p.text)
                            if new != p.text:
                                replace_paragraph_text(p, new)
                                n += 1
        if n:
            doc.save(str(f))
            changed.append((f.name, n))

    print(f"Updated {len(changed)} files")
    for name, n in changed:
        print(f"  {n:3} paras  {name}")

    print("--- verify active ---")
    active = [
        "Syllabus Summer 2026.docx",
        "Psych 4213 - Statistics - Fall 2026.docx",
        "Psych 4213 - Statistics - Fall 2026 (In Person).docx",
        "Psych 4223 - Research Methodology - Fall 2026.docx",
        "Psych 4213 - Statistics - Spring 2026.docx",
    ]
    for name in active:
        p = downloads / name
        if not p.exists():
            print("missing", name)
            continue
        doc = Document(str(p))
        joined = "\n".join(x.text for x in doc.paragraphs)
        print(
            name,
            "Blackboard?",
            bool(re.search(r"[Bb]lack[Bb]oard", joined)),
            "mangled?",
            "Canvashttps" in joined or "blackboardhttps" in joined.lower(),
        )


if __name__ == "__main__":
    main()
